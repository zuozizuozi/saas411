"""Append the post-fix acceptance checkpoint to the single cumulative DOCX."""

import os
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path(os.environ["SUMMARY_DOC"])
OUTPUT = Path(os.environ["SUMMARY_OUTPUT_DOC"])
SECTION_TITLE = "十二、修复后最终验收检查点（2026-08-05）"
NAVY = RGBColor(0x1F, 0x4D, 0x78)
RED = RGBColor(0xB9, 0x1C, 0x1C)
GREEN = RGBColor(0x1C, 0x6B, 0x3C)


def style_run(run, *, size=10.5, bold=False, color=None):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=2):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(5)
    style_run(
        paragraph.add_run(text),
        size=15 if level == 1 else 12,
        bold=True,
        color=NAVY,
    )
    return paragraph


def add_body(doc, lead, body, *, color=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15
    style_run(paragraph.add_run(lead), bold=True, color=color or NAVY)
    style_run(paragraph.add_run(body))
    return paragraph


def add_bullet(doc, text, *, color=None):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.1
    style_run(paragraph.add_run(text), color=color)
    return paragraph


def remove_existing_section(doc):
    body = doc._element.body
    removing = False
    previous = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        text = "".join(child.itertext())
        if SECTION_TITLE in text:
            if previous is not None and previous.tag == qn("w:p"):
                previous_text = "".join(previous.itertext()).strip()
                has_page_break = any(
                    node.get(qn("w:type")) == "page"
                    for node in previous.iter(qn("w:br"))
                )
                if not previous_text and has_page_break:
                    body.remove(previous)
            removing = True
        if removing:
            body.remove(child)
        else:
            previous = child


doc = Document(SOURCE)
remove_existing_section(doc)

add_heading(doc, SECTION_TITLE, level=1)

add_body(
    doc,
    "当前验收状态：",
    "SHIP。高危修复、全量自动化、普通账号/管理员鉴权、一次真实积分生成、生产构建、正式域名切流和上线后日志观察均已通过。",
    color=GREEN,
)

add_heading(doc, "1. 已完成并通过")
for item in [
    "SEC-011 已修复：积分页 returnTo 现在通过 getSafeAuthCallbackURL 只允许站内路径，不再重复 decodeURIComponent。",
    "回归测试完成红绿验证：旧代码上安全测试失败，修复后定向测试 8/8 通过。",
    "全量测试 34 个文件、213/213 通过；TypeScript、Biome、生产构建、依赖审计、机械校验和 git diff 检查均通过。",
    "生产构建生成 182 个静态页面；本地 sitemap 构建阶段生成 84 个 URL。",
    "普通测试账号会话有效，访问 /en/admin 被重定向到公开首页；新标签页继续保持登录状态。",
    "真实生成已完成：Seedance 2.0 Mini、480P、单输出、音频关闭，预计与实际结算均为 35 积分；余额 34201→34166。",
    "生成接口与 Evolink 回调均返回 200，新作品状态 Completed、时长 5s，作品列表和缩略图均可见。",
    "Edge 扩展文件 URL 权限已开启；生产站图生视频页成功选择 icon-192.png，文件名、类型、大小和界面状态校验均正常。",
    "代码确认输入图片采用鉴权后的 presign→对象存储 PUT→complete 验真登记流程；complete 会校验预约所有者、文件元数据、对象大小、Content-Type 和真实图片内容。",
    "真实生成的成品已成功回传并保存到 R2，证明生产环境对象存储写入与公开读取链路正常。",
    "管理员验收通过：上线前后均能访问 /admin；控制台显示 3 个用户、6 个视频、3 个完成、3 个失败、0 个处理中。",
    "生产库只读复核：2 个管理员、1 个普通用户、0 个生成暂停用户、0 个活动任务、0 个超过 10 分钟的卡住任务，因此未执行任务恢复。",
    "Supabase 安全顾问没有 ERROR/WARN；仅有 RLS 已开启但无客户端策略的 INFO，符合当前仅服务端直连数据库的架构。",
]:
    add_bullet(doc, item, color=GREEN if item.startswith("SEC-011") else None)

add_heading(doc, "2. Vercel 复验结果")
for item in [
    "用户已将 BETTER_AUTH_SECRET、NEXT_PUBLIC_APP_URL、GOOGLE_CLIENT_ID、GOOGLE_CLIENT_SECRET、RESEND_API_KEY、RESEND_FROM 配置为 Production and Preview。",
    "用户触发的成功重新部署 dpl_9wDquG52sCGnVZVBNXTWuhKoVN2u 是 production 旧提交 53ab7c0 的重新部署，不包含本地高危修复。",
    "从当前工作区创建的新 Preview dpl_HekzvUW28PpVPHfMdv2DxXg4t6Bf 已达到 READY，远端构建完成 182 个页面并生成 98 个 sitemap URL。",
    "Preview 的 sitemap.xml 返回 200 并包含安全响应头；动态接口受 Vercel Preview SSO 保护，未绕过保护模拟业务登录。",
    "真实生成期间生产日志确认 /api/v1/video/generate 与 /api/v1/video/callback/evolink 均返回 200；近一小时仅有旧生产版本 Credits.packages.team 文案缺失告警，该问题已在候选版本修复。",
    "生产站未被本次 Preview 尝试替换或中断。",
    "新增 .vercelignore，生产上传包不再包含 quality 审计产物、DOCX、渲染图片、本地代理状态、临时文件和维护脚本。",
    "生产隔离部署 dpl_4vCD9k7WLk8m7ivs75mujR4tyTRs 使用 Production 环境变量完成 182/182 页面构建，sitemap 生成 98 个 URL，状态 READY。",
    "隔离部署首页、登录页、文生视频与图生视频页面均在 Edge 通过；随后已提升为 production，并确认 seedance.co.com 正常提供新版本。",
    "正式站 /api/health 与 /api/v1/config/models 返回 200；未登录积分接口返回 401、管理员接口返回 403、任务恢复接口返回 401。",
    "正式站 CSP、HSTS、nosniff、DENY frame、Referrer-Policy 和 Permissions-Policy 响应头均存在。",
    "新部署观察窗口内无 5xx、无 warning/error；旧部署的一次 Better Auth rateLimit 清理超时不属于当前版本。",
]:
    add_bullet(doc, item)

add_heading(doc, "3. 用户目前无需立刻操作")
for item in [
    "网站已经上线并可正常使用；保持 Vercel、Supabase、R2、Evolink 和邮件服务的现有环境变量不变。",
    "源码已在本地分支 agent/security-release-hardening 提交为 dca3f46；GitHub 旧凭据失效导致远端推送未完成，方便时重新登录 GitHub 后再推送该分支即可，不影响现网。",
    "不要随意重放 Stripe/Creem webhook、AI callback 或任务恢复接口；这些入口具备幂等/鉴权保护，但生产环境仍应按真实事件驱动。",
]:
    add_bullet(doc, item)

add_heading(doc, "4. 上线后的人工抽查边界")
for item in [
    "输入图片的生产真实上传与素材入库未单独触发：当前前端只会在提交图生视频时上传；为遵守仅一次真实积分生成的授权，本轮不再发起第二次生成。该项不否定已通过的文件选择、安全实现审计和 R2 成品存储验证。",
    "本轮没有使用真实银行卡发起支付；支付 webhook、幂等、撤销和积分账本已完成代码与自动化检查，真实支付应由首笔正常客户订单自然验证。",
    "Provider Health 最近 60 分钟没有事件样本，当前返回空列表而不是降级；后续产生新任务后再观察 attempts/failures/latency 数据。",
]:
    add_bullet(doc, item)

doc.save(OUTPUT)

check = Document(OUTPUT)
all_text = "\n".join(p.text for p in check.paragraphs)
for required in [
    SECTION_TITLE,
    "SEC-011 已修复",
    "213/213",
    "GOOGLE_CLIENT_ID",
    "dpl_HekzvUW28PpVPHfMdv2DxXg4t6Bf",
]:
    if required not in all_text:
        raise AssertionError(f"Missing acceptance checkpoint content: {required}")
if all_text.count(SECTION_TITLE) != 1:
    raise AssertionError("Acceptance checkpoint section is not idempotent")

print(
    f"DOCX_ACCEPTANCE_CHECKPOINT=PASS PARAGRAPHS={len(check.paragraphs)} "
    f"TABLES={len(check.tables)}"
)
