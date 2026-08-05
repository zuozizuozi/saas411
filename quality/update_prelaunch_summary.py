"""Append the final pre-launch verification to the single cumulative report."""

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(os.environ["SUMMARY_DOC"])
OUTPUT = Path(os.environ["SUMMARY_OUTPUT_DOC"])
SECTION_TITLE = "十一、上线前最终全链路检查（2026-08-05）"
NAVY = RGBColor(0x1F, 0x4D, 0x78)
RED = RGBColor(0xB9, 0x1C, 0x1C)
GREEN = RGBColor(0x1C, 0x6B, 0x3C)
GRAY = RGBColor(0x4B, 0x55, 0x63)


def style_run(run, *, size=10.5, bold=False, color=None):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_paragraph(paragraph, *, after=5, line=1.15):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    for run in paragraph.runs:
        style_run(run)


def add_body(doc, text, *, lead=None, color=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15
    if lead and text.startswith(lead):
        style_run(paragraph.add_run(lead), bold=True, color=color or NAVY)
        style_run(paragraph.add_run(text[len(lead) :]))
    else:
        style_run(paragraph.add_run(text), color=color)
    return paragraph


def add_bullet(doc, text, *, bold_lead=None):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_lead and text.startswith(bold_lead):
        style_run(paragraph.add_run(bold_lead), bold=True, color=NAVY)
        style_run(paragraph.add_run(text[len(bold_lead) :]))
    else:
        style_run(paragraph.add_run(text))
    return paragraph


def set_cell(cell, text, *, header=False, status=None):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    color = NAVY if header else GREEN if status == "PASS" else RED if status == "FAIL" else None
    style_run(paragraph.add_run(text), size=8.8, bold=header or status in {"PASS", "FAIL"}, color=color)


def remove_existing_final_section(doc):
    body = doc._element.body
    removing = False
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        text = "".join(child.itertext())
        if SECTION_TITLE in text:
            removing = True
        if removing:
            body.remove(child)


doc = Document(SOURCE)
remove_existing_final_section(doc)

page_break = doc.add_paragraph()
page_break.add_run().add_break(WD_BREAK.PAGE)

heading = doc.add_paragraph(style="Heading 1")
heading.paragraph_format.keep_with_next = True
heading.paragraph_format.space_after = Pt(8)
style_run(heading.add_run(SECTION_TITLE), size=15, bold=True, color=NAVY)

verdict = doc.add_paragraph()
verdict.paragraph_format.space_after = Pt(8)
verdict.paragraph_format.line_spacing = 1.15
style_run(verdict.add_run("发布裁决：FIX BEFORE MERGE"), size=12, bold=True, color=RED)
style_run(
    verdict.add_run(
        "。代码门禁通过，但存在 1 个高危 DOM XSS、最新预览部署失败，且统一修复尚未部署到当前生产版本。"
    )
)

add_body(
    doc,
    "检查范围：公开页面、登录入口、普通用户/管理员 API、AI/Stripe/QStash 回调、Cron、模型目录、积分与视频链路、存储入口、环境变量、线上日志、安全响应头及桌面/移动端浏览器行为。",
    lead="检查范围：",
)
add_body(
    doc,
    "版本边界：工作区分支 agent/security-release-hardening，HEAD e781082，包含尚未提交的统一修复；当前生产仍为 main 的 53ab7c0。",
    lead="版本边界：",
)

for title in ["1. 必须在合并前解决", "2. 自动化与线上门禁证据"]:
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(5)
    style_run(p.add_run(title), size=12, bold=True, color=NAVY)
    if title.startswith("1."):
        add_bullet(
            doc,
            "SEC-011（高危，高置信度）：src/components/credits/credits-page.tsx 第 62-83 行把 returnTo 查询参数经二次 decode 后直接传给 router.push。登录用户访问攻击链接时可形成开放跳转，并可能执行 javascript: URL。必须改用现有 getSafeAuthCallbackURL 白名单校验并补负向测试。",
            bold_lead="SEC-011（高危，高置信度）：",
        )
        add_bullet(
            doc,
            "REL-011（发布阻断）：最新分支预览部署 dpl_Df8z9VWTqrCE9DULxeM3UYeiAit3 为 ERROR；构建日志确认 Preview 环境缺少 BETTER_AUTH_SECRET 与 NEXT_PUBLIC_APP_URL。需补齐 Preview 作用域变量并获得 READY 预览。",
            bold_lead="REL-011（发布阻断）：",
        )
        add_bullet(
            doc,
            "REL-012（版本阻断）：8 项统一修复当前只在本地工作区，尚未提交、未产生绿色预览、未进入生产；线上仍可复现旧版翻译缺失。",
            bold_lead="REL-012（版本阻断）：",
        )

table = doc.add_table(rows=1, cols=3)
table.autofit = False
table.style = "Table Grid"
widths = [Inches(1.45), Inches(0.85), Inches(4.0)]
headers = ["门禁", "结果", "证据"]
for idx, text in enumerate(headers):
    table.columns[idx].width = widths[idx]
    set_cell(table.cell(0, idx), text, header=True)

rows = [
    ("TypeScript", "PASS", "tsc --noEmit，无错误"),
    ("Biome", "PASS", "395 个文件，无错误"),
    ("全量测试", "PASS", "33 个测试文件，211/211 通过"),
    ("生产构建", "PASS", "Node 24 + Next.js 15.5.21；SKIP_ENV_VALIDATION=1 时完成 182 个静态页面"),
    ("依赖漏洞", "PASS", "pnpm audit --prod --audit-level high：无已知漏洞"),
    ("密钥扫描", "PASS", "当前文件、受控敏感文件与 Git 历史均无真实高置信度凭据命中"),
    ("鉴权负测", "PASS", "普通接口 401；管理员接口 401/403；Cron、AI、Stripe、QStash 均拒绝无签名请求"),
    ("OAuth 初始化", "PASS", "同源 Google 登录请求 200，回调为 /api/auth/callback/google；缺失 Origin 请求被 403 拒绝"),
    ("浏览器冒烟", "PASS", "首页、定价、登录、注册、工具页均 200；未登录生成/购买不触发付费调用；390px 无横向溢出"),
    ("机械复核", "PASS", "Provider 机械制品与源码一致"),
    ("Vercel 预览", "FAIL", "最新分支部署因 Preview 环境变量缺失而失败"),
]
for gate, status, evidence in rows:
    cells = table.add_row().cells
    for idx, value in enumerate((gate, status, evidence)):
        cells[idx].width = widths[idx]
        set_cell(cells[idx], value, status=status if idx == 1 else None)

p = doc.add_paragraph(style="Heading 2")
p.paragraph_format.keep_with_next = True
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(5)
style_run(p.add_run("3. 已确认的非核心线上问题"), size=12, bold=True, color=NAVY)
add_bullet(doc, "当前生产文本/图生视频页仍报 Credits.packages.team.name/description 缺失；修复后本地生产包中英文均已无该错误，部署新版本后应复验。")
add_bullet(doc, "未登录首页、定价页和工具页会请求受保护的积分/素材接口并产生预期 401 控制台错误；不阻断使用，但会污染监控。")
add_bullet(doc, "线上 sitemap.xml 返回 200 但 URL 数为 0；运行时扫描 /var/task/src/app/[locale] 失败。属于 SEO/发布完整性缺陷，不影响核心生成链路。")
add_bullet(doc, "无签名 video-reconcile 请求被 QStash 正确拦截，但主入口返回 500 而非 401/403，可能造成误报警。")

p = doc.add_paragraph(style="Heading 2")
p.paragraph_format.keep_with_next = True
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(5)
style_run(p.add_run("4. 今天上线前必须人工完成"), size=12, bold=True, color=NAVY)
manual = [
    "修复 SEC-011 并添加 javascript:、https://、//evil.example、畸形百分号四类负向测试；重新跑 211 项测试、类型、Lint 与构建。",
    "在 Vercel Preview 作用域补齐所需变量，确认预览部署 READY；再核对 Production 变量名称与作用域，不展示或复制真实值。",
    "用专用测试账号分别完成 Google 登录、邮箱 OTP、退出与重新登录；普通账号访问 /admin 必须拒绝。",
    "上传一张小尺寸真实图片，确认 presign → PUT → complete → 素材列表；用第二账号确认不能引用第一账号图片。",
    "执行一次最低成本 Evolink 生成，观察积分冻结、任务 ID、回调或 QStash、R2 入库、COMPLETED 与积分结算；再做一次可控失败并确认释放积分。",
    "仅在 Stripe 测试模式完成一次订阅与一次积分包支付，确认 webhook 幂等、订单/积分账本、支付返回页；不得用真实银行卡做上线验证。",
    "管理员检查 Provider Health、用户/角色、视频恢复入口；发布后 15 分钟检查 Vercel 5xx、回调、数据库连接、QStash 与存储日志。",
    "确认生产别名 seedance.co.com 指向本次目标 SHA，并复测首页、登录、生成、积分、作品列表与安全响应头。",
]
for item in manual:
    add_bullet(doc, item)

p = doc.add_paragraph(style="Heading 2")
p.paragraph_format.keep_with_next = True
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(5)
style_run(p.add_run("5. 安全审查结论"), size=12, bold=True, color=NAVY)
add_body(doc, "安全发现汇总：严重 0，高危 1，中危 0，低危 0；生产依赖漏洞 0；真实凭据暴露 0。", lead="安全发现汇总：")
add_body(
    doc,
    "已验证安全边界：SQL 均通过 Drizzle 参数化；用户资源查询携带 userId；管理员角色实时查库；上传绑定预约与所有权；回调/定时任务需要签名或密钥；OAuth 缺失同源 Origin 被拒绝。",
    lead="已验证安全边界：",
)

doc.save(OUTPUT)

check = Document(OUTPUT)
all_text = "\n".join(p.text for p in check.paragraphs)
required = [SECTION_TITLE, "FIX BEFORE MERGE", "SEC-011", "211/211", "今天上线前必须人工完成"]
for needle in required:
    if needle not in all_text and not any(
        needle in cell.text
        for table_item in check.tables
        for row in table_item.rows
        for cell in row.cells
    ):
        raise AssertionError(f"Missing final pre-launch content: {needle}")
if all_text.count(SECTION_TITLE) != 1:
    raise AssertionError("Final pre-launch section is not idempotent")
if len(check.tables) != 4:
    raise AssertionError(f"Unexpected table count: {len(check.tables)}")

print(
    f"DOCX_PRELAUNCH=PASS PARAGRAPHS={len(check.paragraphs)} "
    f"TABLES={len(check.tables)}"
)
