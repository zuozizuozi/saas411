"""Update the single cumulative desktop DOCX with the unified repair result."""

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DOCX = Path(os.environ["SUMMARY_DOC"])
OUTPUT = Path(os.environ["SUMMARY_OUTPUT_DOC"])
NAVY = RGBColor(0x1F, 0x4D, 0x78)


def style_run(run, size=10.5, bold=False, color=None):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def replace_paragraph(paragraph, text, *, bold_lead=None):
    paragraph.clear()
    if bold_lead and text.startswith(bold_lead):
        style_run(paragraph.add_run(bold_lead), bold=True, color=NAVY)
        style_run(paragraph.add_run(text[len(bold_lead) :]))
    else:
        style_run(paragraph.add_run(text))


def set_cell(cell, text, *, bold=False, size=9):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        for run in paragraph.runs:
            style_run(run, size=size, bold=bold)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    for existing in tr_pr.findall(qn("w:tblHeader")):
        tr_pr.remove(existing)
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


doc = Document(DOCX)

replacements = {
    "阶段总结（累计更新至第 6 阶段，全部完成）": "阶段总结（6 阶段审计 + 统一修复，累计完成）",
    "阶段性判断：全量审计 6 个阶段均已完成；8 个有效问题全部复现，审计制品最终门禁通过，但产品仍不具备发布条件。": "阶段性判断：全量审计 6 个阶段和统一修复均已完成；8 个有效问题全部关闭，当前无仍处于打开状态的已确认源码缺陷。",
    "证据强度：27 项功能测试、完整 200 项测试、生产构建、类型检查、Lint、机械复核、19 个补丁检查和最终质量门禁均有运行凭证。": "证据强度：27 项功能测试、211 项全量测试、生产构建、类型检查、Lint、依赖审计、密钥扫描、机械复核和最终质量门禁均有运行凭证。",
    "最终结果：质量门禁 93 条通过记录、0 项失败、3 条旧清单兼容提醒；8 个问题中 6 个拟议修复通过 GREEN，1 个不完整，1 个无获批补丁。": "最终结果：8/8 个确认问题已修复并复核；全量测试 211/211 通过，质量门禁 0 项失败、3 条旧清单兼容提醒。",
    "源码状态：所有红绿验证均在一次性隔离工作区完成并清理；主工作区仅新增 quality 审计制品，没有修改网站产品源码。": "第 5 阶段当时源码状态：所有红绿验证均在一次性隔离工作区完成并清理；当时主工作区尚未修改网站产品源码。",
    "发布判断：审计流程已完成，但发布建议仍为 BLOCK：3 个高风险问题尚未进入产品源码，BUG-007 无获批修复，BUG-008 补丁仍不完整。": "第 6 阶段当时发布判断：审计已完成，但 8 个问题尚未统一进入产品源码，因此当时建议保持 BLOCK。",
    "源码状态：主工作区没有产品源码修改；所有变更仅位于 quality 审计目录和本累计总结文档。": "第 6 阶段当时源码状态：产品源码尚未修改；统一修复在用户后续明确授权后执行。",
    "九、后续处置建议": "九、统一修复：8 个问题全部落地",
    "处置原则：审计已经完成；后续工作是按风险批次实施修复并复核，不把未应用的拟议补丁当作已修复。": "执行方式：在用户明确授权后一次性修改产品源码，并保留原始 RED 证据；所有活动回归已由预期失败改为普通通过测试。",
    "优先处理 BUG-002、BUG-007、BUG-009 三个高风险问题，并由负责人确认长任务超时和 DNS 连接固定方案。": "高风险闭环：BUG-002 增加 60 分钟终态与上传租约竞态保护；BUG-007 将实际连接固定到已验证公网 IP；BUG-009 统一 QStash 发送端、接收端和失败回调的完整配置要求。",
    "随后处理 BUG-004、BUG-005、BUG-008、BUG-010 与 BUG-006；BUG-008 必须补齐 payment_reversal 的全部语言翻译。": "其余闭环：回调请求体限制为 1 MiB；修正文档密钥名和 Vercel CSP；统一积分流水词汇并补齐 7 种语言；贯通作品状态、模型、排序与数据库分页。",
    "修复应分批应用，每批运行对应回归测试、完整测试、类型检查、Lint、生产构建和质量门禁。": "验证结果：211/211 测试、TypeScript、Biome、Next.js 生产构建、依赖漏洞审计、已跟踪文件密钥扫描、机械验证和 Quality Playbook 门禁全部通过。",
    "真实 Provider、支付、数据库和存储链路必须在获批的可清理测试环境中验证；缺少授权时继续保持 BLOCK。": "剩余部署门禁：本轮未调用真实付费 Provider、支付和对象存储；上线前仍需在可清理测试环境注入完整环境变量并执行一次真实冒烟测试。",
    "十、审计文件位置": "十、审计与统一修复文件位置",
}

for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    replacement = replacements.get(text)
    if replacement:
        replace_paragraph(paragraph, replacement)
    if paragraph.text.strip() in {
        "九、统一修复：8 个问题全部落地",
        "十、审计与统一修复文件位置",
    }:
        paragraph.paragraph_format.keep_with_next = True

# Overview table.
overview = doc.tables[0]
set_cell(overview.cell(1, 1), "6 个审计阶段与统一修复均已完成；进入部署前真实环境冒烟门禁")
set_cell(overview.cell(2, 1), "8 个有效问题已全部修复：3 个高、4 个中、1 个低；当前活动确认问题为 0")
set_cell(overview.cell(3, 1), "注入完整生产环境变量，在可清理环境验证真实 Provider、支付、数据库与存储")

# Current issue status table. Historical severity stays visible; status becomes current.
issues = doc.tables[1]
repeat_table_header(issues.rows[0])
issue_updates = {
    2: ("已修复：60 分钟终态策略，并保护进行中的上传租约与积分结算。", "REQ-007 / BUG-002 / FIXED"),
    4: ("已修复：声明长度和实际流量均受 1 MiB 上限约束，异常 JSON 返回 400。", "REQ-006 / BUG-004 / FIXED"),
    5: ("已修复：活动指南、示例环境和可执行代码统一使用 CALLBACK_HMAC_SECRET。", "REQ-012 / BUG-005 / FIXED"),
    6: ("已修复：CSP 精确允许已启用的 Vercel 观测脚本源。", "REQ-014 / BUG-006 / FIXED"),
    7: ("已修复：每次连接和重定向均使用已验证公网 IP，并保留 Host/SNI、代理、限额和清理。", "REQ-010 / BUG-007 / FIXED"),
    8: ("已修复：API、类型、UI 和 7 种语言使用同一完整词汇，含 payment_reversal。", "REQ-013 / BUG-008 / FIXED"),
    9: ("已修复：发送、接收和失败回调共享完整 QStash 配置契约，部分配置不再误调度。", "REQ-007、012 / BUG-009 / FIXED"),
    10: ("已修复：状态、模型、排序进入数据库查询；升序分页方向同步修正。", "REQ-013 / BUG-010 / FIXED"),
}
for row_index, (impact, tracking) in issue_updates.items():
    set_cell(issues.cell(row_index, 2), impact)
    set_cell(issues.cell(row_index, 3), tracking)

# Keep the final verification snapshot in the compact stage-2 table current.
verification = doc.tables[2]
set_cell(verification.cell(1, 2), "27/27；全量测试 211/211")
set_cell(verification.cell(2, 2), "无类型错误")
set_cell(verification.cell(3, 2), "395 个文件，无错误")
set_cell(verification.cell(4, 2), "通过；Provider 制品与源码一致")
set_cell(verification.cell(5, 1), "通过")
set_cell(verification.cell(5, 2), "质量门禁 0 FAIL、3 个旧清单兼容 WARN")

doc.save(OUTPUT)

# Structural post-save checks before the file replaces the desktop copy.
check = Document(OUTPUT)
text = "\n".join(p.text for p in check.paragraphs)
required = [
    "6 阶段审计 + 统一修复",
    "九、统一修复：8 个问题全部落地",
    "211/211",
    "当前活动确认问题为 0",
    "真实付费 Provider、支付和对象存储",
]
for needle in required:
    if needle not in text and not any(
        needle in cell.text for table in check.tables for row in table.rows for cell in row.cells
    ):
        raise AssertionError(f"Missing unified repair summary text: {needle}")
if text.count("九、统一修复：8 个问题全部落地") != 1:
    raise AssertionError("Unified repair heading count is not one")
if len(check.tables) != 3:
    raise AssertionError(f"Unexpected table count: {len(check.tables)}")

print(
    f"DOCX_UNIFIED_REPAIR=PASS PARAGRAPHS={len(check.paragraphs)} "
    f"TABLES={len(check.tables)}"
)
